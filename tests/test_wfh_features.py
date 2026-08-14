import sqlite3
import tempfile
import unittest
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

from drive_storage import safe_drive_name
from storage import Storage
from wfh_report import (EvidenceFile, ReportActivity, build_wfh_report,
                        report_date_for_access, report_generation_allowed,
                        report_window_open, tidy_sentence)


JAKARTA = ZoneInfo("Asia/Jakarta")
ROOT = Path(__file__).resolve().parents[1]


class FridayWindowTests(unittest.TestCase):
    def test_window_is_open_until_last_second_of_friday(self):
        self.assertTrue(report_window_open(datetime(2026, 8, 14, 23, 59, 59,
                                                    tzinfo=JAKARTA)))
        self.assertFalse(report_window_open(datetime(2026, 8, 15, 0, 0, 0,
                                                     tzinfo=JAKARTA)))
        self.assertFalse(report_window_open(datetime(2026, 8, 13, 23, 59, 59,
                                                     tzinfo=JAKARTA)))

    def test_tidy_sentence_does_not_need_ai(self):
        self.assertEqual("Mengikuti rapat internal.",
                         tidy_sentence("  mengikuti   rapat internal "))

    def test_main_never_uses_naive_railway_server_clock(self):
        source = (ROOT / "main.py").read_text(encoding="utf-8")
        self.assertNotIn("datetime.now()", source)

    def test_admin_can_override_friday_window(self):
        saturday = datetime(2026, 8, 15, 9, 0, tzinfo=JAKARTA)
        friday = datetime(2026, 8, 14, 9, 0, tzinfo=JAKARTA)
        self.assertTrue(report_generation_allowed(saturday, "open"))
        self.assertFalse(report_generation_allowed(friday, "closed"))
        self.assertFalse(report_generation_allowed(saturday, "auto"))
        self.assertEqual(friday.date(), report_date_for_access(saturday, "open"))


class DailyStorageTests(unittest.TestCase):
    def test_report_access_mode_defaults_to_auto_and_persists(self):
        with tempfile.TemporaryDirectory() as tmp:
            storage = Storage(str(Path(tmp) / "settings.db"))
            self.assertEqual("auto", storage.wfh_report_access_mode())
            storage.set_wfh_report_access_mode("open")
            self.assertEqual("open", storage.wfh_report_access_mode())

    def test_non_asn_is_active_without_password(self):
        with tempfile.TemporaryDirectory() as tmp:
            storage = Storage(str(Path(tmp) / "non-asn.db"))
            storage.invite_user(2001, "NIK-001", "PEGAWAI UJI",
                                employee_type="non_asn")
            user = storage.get_user(2001)
        self.assertEqual("active", user[4])
        self.assertIsNone(user[2])
        self.assertEqual("non_asn", user[8])

    def test_activity_insert_is_idempotent_per_telegram_message(self):
        with tempfile.TemporaryDirectory() as tmp:
            storage = Storage(str(Path(tmp) / "idempotent.db"))
            kwargs = dict(
                telegram_id=2001, employee_type="non_asn",
                activity_date="2026-08-12", activity_time="10:00:00",
                activity_text="Mengikuti rapat internal.",
                created_at_local="2026-08-12T10:00:00+07:00",
                source_chat_id=2001, source_message_id=88,
            )
            first = storage.add_daily_activity(**kwargs)
            second = storage.add_daily_activity(**kwargs)
        self.assertEqual(first[0], second[0])
        self.assertTrue(first[1])
        self.assertFalse(second[1])

    def test_wfh_query_rejects_activity_entered_after_friday(self):
        with tempfile.TemporaryDirectory() as tmp:
            storage = Storage(str(Path(tmp) / "window.db"))
            common = dict(
                telegram_id=2001, employee_type="non_asn",
                activity_date="2026-08-14", activity_time="10:00:00",
                activity_text="Aktivitas WFH.", source_chat_id=2001,
            )
            storage.add_daily_activity(
                **common, created_at_local="2026-08-14T10:00:00+07:00",
                source_message_id=1)
            storage.add_daily_activity(
                **common, created_at_local="2026-08-15T08:00:00+07:00",
                source_message_id=2)
            rows = storage.list_wfh_activities(2001, "2026-08-14")
        self.assertEqual(1, len(rows))

    def test_old_users_database_migrates_without_losing_account(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "legacy.db"
            db = sqlite3.connect(path)
            db.execute("""CREATE TABLE users (
              telegram_id INTEGER PRIMARY KEY,nip TEXT NOT NULL,password_enc TEXT,
              full_name TEXT,status TEXT NOT NULL DEFAULT 'invited',
              is_admin INTEGER NOT NULL DEFAULT 0,created_at TEXT DEFAULT CURRENT_TIMESTAMP)""")
            db.execute("""INSERT INTO users(telegram_id,nip,full_name,status)
              VALUES(1001,'199001012020121001','PEGAWAI LAMA','active')""")
            db.commit()
            db.close()
            storage = Storage(str(path))
            user = storage.get_user(1001)
        self.assertEqual("PEGAWAI LAMA", user[3])
        self.assertEqual("asn", user[8])
        self.assertEqual("Pemasaran", user[9])


class WfhDocumentTests(unittest.TestCase):
    def test_report_preserves_photo_and_required_alignment(self):
        with tempfile.TemporaryDirectory() as tmp:
            temp = Path(tmp)
            photo = temp / "photo.jpg"
            Image.new("RGB", (1600, 900), color=(40, 90, 130)).save(photo)
            output = temp / "report.docx"
            build_wfh_report(
                template_path=ROOT / "templates" / "Laporan_WFH_template.docx",
                output_path=output,
                employee_name="PEGAWAI UJI", employee_identifier="199000000000000001",
                employee_type="asn", position="Jabatan Uji", unit_name="Pemasaran",
                report_date=datetime(2026, 8, 14).date(),
                activities=[ReportActivity(
                    activity_time="09:30:00",
                    text="Mengikuti rapat koordinasi dan menyusun tindak lanjut kegiatan.",
                    evidence=[
                        EvidenceFile("https://drive.google.com/file/d/test-1/view", photo),
                        EvidenceFile("https://drive.google.com/file/d/test-2/view", photo),
                    ],
                )],
            )
            document = Document(output)
            activity_table = document.tables[1]
            date_paragraph = activity_table.cell(1, 1).paragraphs[0]
            activity_paragraph = activity_table.cell(1, 2).paragraphs[0]
            relationships = list(document.part.rels.values())
            fonts = [run.font.name for table in document.tables for row in table.rows
                     for cell in row.cells for paragraph in cell.paragraphs
                     for run in paragraph.runs if run.text]
        self.assertEqual(WD_ALIGN_PARAGRAPH.LEFT, date_paragraph.alignment)
        self.assertEqual(WD_ALIGN_PARAGRAPH.JUSTIFY, activity_paragraph.alignment)
        self.assertEqual(2, len(document.inline_shapes))
        self.assertIn("Arial", fonts)
        self.assertTrue(any(getattr(rel, "target_ref", "").startswith("https://drive.google.com/")
                            for rel in relationships))

    def test_drive_names_cannot_create_unintended_paths(self):
        self.assertEqual("Nama Pegawai Bukti", safe_drive_name("Nama/ Pegawai\\Bukti"))

    def test_non_asn_report_uses_employee_id_instead_of_nip(self):
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "non-asn.docx"
            build_wfh_report(
                template_path=ROOT / "templates" / "Laporan_WFH_template.docx",
                output_path=output,
                employee_name="PEGAWAI NON ASN", employee_identifier="NIK-001",
                employee_type="non_asn", position="Tenaga Pendukung",
                unit_name="Pemasaran", report_date=datetime(2026, 8, 14).date(),
                activities=[ReportActivity("10:00:00", "Mengikuti rapat internal.")],
            )
            document = Document(output)
            profile = document.tables[0]
            signature = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertEqual("ID Pegawai", profile.cell(1, 0).text)
        self.assertEqual("NIK-001", profile.cell(1, 2).text)
        self.assertIn("ID Pegawai. NIK-001", signature)


if __name__ == "__main__":
    unittest.main()
