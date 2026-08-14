from __future__ import annotations

import os
import re
import shutil
import subprocess
import io
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt
from PIL import Image


JAKARTA = ZoneInfo("Asia/Jakarta")
DAY_NAMES = ("Senin", "Selasa", "Rabu", "Kamis", "Jumat", "Sabtu", "Minggu")
MONTH_NAMES = (
    "", "Januari", "Februari", "Maret", "April", "Mei", "Juni",
    "Juli", "Agustus", "September", "Oktober", "November", "Desember",
)


@dataclass(frozen=True)
class EvidenceFile:
    url: str
    local_path: Path | None = None
    file_name: str = ""


@dataclass(frozen=True)
class ReportActivity:
    activity_time: str
    text: str
    evidence: list[EvidenceFile] = field(default_factory=list)


def now_jakarta() -> datetime:
    return datetime.now(JAKARTA)


def report_window_open(moment: datetime | None = None) -> bool:
    moment = moment or now_jakarta()
    if moment.tzinfo is None:
        moment = moment.replace(tzinfo=JAKARTA)
    else:
        moment = moment.astimezone(JAKARTA)
    return moment.weekday() == 4  # Jumat, 00.00.00 sampai 23.59.59 WIB.


def report_generation_allowed(moment: datetime, access_mode: str = "auto") -> bool:
    """Terapkan jadwal Jumat atau override buka/tutup dari admin."""
    if access_mode == "open":
        return True
    if access_mode == "closed":
        return False
    return report_window_open(moment)


def report_date_for_access(moment: datetime, access_mode: str = "auto") -> date:
    """Mode manual tetap mengarah ke Jumat terakhir, bukan tanggal server saat ini."""
    local_moment = moment if moment.tzinfo is None else moment.astimezone(JAKARTA)
    if access_mode == "open":
        days_since_friday = (local_moment.weekday() - 4) % 7
        return (local_moment - timedelta(days=days_since_friday)).date()
    return local_moment.date()


def tidy_sentence(value: str) -> str:
    """Perapian deterministik tanpa AI dan tanpa mengubah isi substantif."""
    text = " ".join((value or "").strip().split())
    if not text:
        return ""
    text = text[0].upper() + text[1:]
    if text[-1] not in ".!?":
        text += "."
    return text


def format_indonesian_date(value: date, include_day: bool = False) -> str:
    base = f"{value.day} {MONTH_NAMES[value.month]} {value.year}"
    return f"{DAY_NAMES[value.weekday()]}, {base}" if include_day else base


def report_file_name(report_date: date) -> str:
    return f"Laporan_WFH_{report_date.day}_{MONTH_NAMES[report_date.month]}_{report_date.year}.docx"


def report_pdf_file_name(report_date: date) -> str:
    return f"Laporan_WFH_{report_date.day}_{MONTH_NAMES[report_date.month]}_{report_date.year}.pdf"


def convert_docx_to_pdf(docx_path: str | Path, output_dir: str | Path) -> Path:
    """Konversi DOCX ke PDF memakai LibreOffice pada direktori sementara."""
    source = Path(docx_path).resolve()
    if not source.is_file():
        raise FileNotFoundError("File Word laporan tidak ditemukan untuk konversi PDF.")

    executable = shutil.which("libreoffice") or shutil.which("soffice")
    if not executable:
        raise RuntimeError(
            "Konverter PDF LibreOffice belum terpasang pada server. "
            "Pastikan konfigurasi Nixpacks ikut dideploy.")

    destination = Path(output_dir).resolve()
    destination.mkdir(parents=True, exist_ok=True)
    profile = destination / ".libreoffice-profile"
    home = destination / ".libreoffice-home"
    profile.mkdir(parents=True, exist_ok=True)
    home.mkdir(parents=True, exist_ok=True)

    environment = os.environ.copy()
    environment["HOME"] = str(home)
    environment["TMPDIR"] = str(destination)
    command = [
        executable,
        "--headless",
        f"-env:UserInstallation={profile.as_uri()}",
        "--convert-to",
        "pdf:writer_pdf_Export",
        "--outdir",
        str(destination),
        str(source),
    ]
    try:
        result = subprocess.run(
            command,
            check=False,
            capture_output=True,
            text=True,
            timeout=120,
            env=environment,
        )
    except subprocess.TimeoutExpired as exc:
        raise RuntimeError("Konversi PDF melewati batas waktu 120 detik.") from exc

    output = destination / f"{source.stem}.pdf"
    if result.returncode != 0 or not output.is_file() or output.stat().st_size == 0:
        raise RuntimeError(
            f"Konversi PDF gagal (kode proses {result.returncode}).")
    return output


def _set_run_font(run, *, bold: bool | None = None, color: str | None = None,
                  underline: bool | None = None):
    run.font.name = "Arial"
    run.font.size = Pt(12)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = None
        color_el = run._element.get_or_add_rPr().find(qn("w:color"))
        if color_el is None:
            color_el = OxmlElement("w:color")
            run._element.get_or_add_rPr().append(color_el)
        color_el.set(qn("w:val"), color)
    if underline is not None:
        run.underline = underline
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), "Arial")


def _format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    for run in paragraph.runs:
        _set_run_font(run)


def _clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _set_cell_text(cell, text: str, alignment=WD_ALIGN_PARAGRAPH.CENTER, *, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text)
    _set_run_font(run, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_hyperlink(paragraph, text: str, url: str):
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(key), "Arial")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "24")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend((fonts, size, size_cs, color, underline))
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_picture(paragraph, path: Path):
    with Image.open(path) as image:
        width_px, height_px = image.size
    if width_px <= 0 or height_px <= 0:
        return
    ratio = width_px / height_px
    max_width = 1.35
    max_height = 1.05
    width = min(max_width, max_height * ratio)
    height = width / ratio
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width), height=Inches(height))


def _floating_anchor_from_inline(inline, *, horizontal_offset: int,
                                 vertical_offset: int):
    """Ubah gambar inline menjadi anchor di depan teks untuk tanda tangan natural."""
    anchor = OxmlElement("wp:anchor")
    for name, value in {
        "distT": "0", "distB": "0", "distL": "0", "distR": "0",
        "simplePos": "0", "relativeHeight": "251658240", "behindDoc": "0",
        "locked": "0", "layoutInCell": "1", "allowOverlap": "1",
    }.items():
        anchor.set(name, value)

    simple_position = OxmlElement("wp:simplePos")
    simple_position.set("x", "0")
    simple_position.set("y", "0")
    anchor.append(simple_position)

    position_horizontal = OxmlElement("wp:positionH")
    position_horizontal.set("relativeFrom", "column")
    horizontal = OxmlElement("wp:posOffset")
    horizontal.text = str(horizontal_offset)
    position_horizontal.append(horizontal)
    anchor.append(position_horizontal)

    position_vertical = OxmlElement("wp:positionV")
    position_vertical.set("relativeFrom", "paragraph")
    vertical = OxmlElement("wp:posOffset")
    vertical.text = str(vertical_offset)
    position_vertical.append(vertical)
    anchor.append(position_vertical)

    extent = inline.find(qn("wp:extent"))
    effect_extent = inline.find(qn("wp:effectExtent"))
    document_properties = inline.find(qn("wp:docPr"))
    frame_properties = inline.find(qn("wp:cNvGraphicFramePr"))
    graphic = inline.find(qn("a:graphic"))
    if any(item is None for item in (
            extent, document_properties, frame_properties, graphic)):
        raise ValueError("Struktur gambar tanda tangan tidak valid.")
    anchor.append(extent)
    if effect_extent is None:
        effect_extent = OxmlElement("wp:effectExtent")
        for name in ("l", "t", "r", "b"):
            effect_extent.set(name, "0")
    anchor.append(effect_extent)
    anchor.append(OxmlElement("wp:wrapNone"))
    document_properties.set("name", "Tanda Tangan Pegawai")
    document_properties.set("descr", "Tanda tangan pegawai")
    anchor.append(document_properties)
    anchor.append(frame_properties)
    anchor.append(graphic)
    inline.getparent().replace(inline, anchor)


def _add_floating_signature(paragraph, path: Path):
    """Pangkas margin transparan dan pasang tanda tangan mengikuti contoh final."""
    with Image.open(path) as image:
        rgba = image.convert("RGBA")
    alpha_box = rgba.getchannel("A").getbbox()
    if alpha_box is None:
        raise ValueError("PNG tanda tangan tidak memiliki gambar yang terlihat.")
    left, top, right, bottom = alpha_box
    padding = max(8, int(max(right - left, bottom - top) * 0.03))
    cropped = rgba.crop((
        max(0, left - padding), max(0, top - padding),
        min(rgba.width, right + padding), min(rgba.height, bottom + padding),
    ))
    width_px, height_px = cropped.size
    ratio = width_px / height_px
    max_width = 1.174
    max_height = 1.107
    if max_width / ratio <= max_height:
        width, height = max_width, max_width / ratio
    else:
        height, width = max_height, max_height * ratio

    signature_stream = io.BytesIO()
    cropped.save(signature_stream, format="PNG", optimize=True)
    signature_stream.seek(0)
    drawing = paragraph.add_run().add_picture(
        signature_stream, width=Inches(width), height=Inches(height))
    inline = drawing._inline
    _floating_anchor_from_inline(
        inline,
        horizontal_offset=455960,
        vertical_offset=73439,
    )


def _repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    header = properties.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        properties.append(header)
    header.set(qn("w:val"), "true")


def build_wfh_report(*, template_path: str | Path, output_path: str | Path,
                     employee_name: str, employee_identifier: str,
                     employee_type: str, position: str, unit_name: str,
                     report_date: date, activities: list[ReportActivity],
                     signature_path: str | Path | None = None) -> Path:
    template = Path(template_path)
    if not template.is_file():
        raise FileNotFoundError("Template laporan WFH tidak ditemukan.")
    document = Document(template)
    if len(document.tables) < 2:
        raise ValueError("Struktur template laporan WFH tidak valid.")

    identifier_label = "NIP" if employee_type == "asn" else "ID Pegawai"
    profile_table = document.tables[0]
    values = {
        "Nama": employee_name,
        "NIP": employee_identifier if employee_type == "asn" else employee_identifier,
        "Nama Jabatan": position or "-",
        "Unit Kerja": unit_name or "Pemasaran",
        "Tanggal WFH": format_indonesian_date(report_date),
    }
    if employee_type != "asn":
        profile_table.cell(1, 0).text = identifier_label
    for row in profile_table.rows:
        label = row.cells[0].text.strip()
        canonical_label = "NIP" if label in {"NIP", "ID Pegawai"} else label
        if canonical_label in values:
            _set_cell_text(row.cells[0], identifier_label if canonical_label == "NIP" else canonical_label,
                           WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell_text(row.cells[1], ":", WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row.cells[2], values[canonical_label], WD_ALIGN_PARAGRAPH.LEFT)

    signature_anchor_paragraph = None
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if "JEJAK ASN" in text:
            paragraph.text = text.replace("JEJAK ASN", "E-MASTER JATIM")
        elif re.match(r"^Surabaya,\s+", text, flags=re.IGNORECASE):
            paragraph.text = f"Surabaya, {format_indonesian_date(report_date)}"
            signature_anchor_paragraph = paragraph
        elif (text.upper().startswith("NIP.") or text.upper().startswith("ID PEGAWAI.")
              or "[[IDENTIFIER_LABEL]]" in text):
            paragraph.text = f"{identifier_label}. {employee_identifier}"
        elif text.upper() == "ACHMAD RIZA MAULANA, S.T" or text == "[[NAMA_PEGAWAI]]":
            paragraph.text = employee_name
        if paragraph.text.strip():
            _format_paragraph(paragraph, paragraph.alignment)

    if signature_path:
        signature = Path(signature_path)
        if not signature.is_file():
            raise FileNotFoundError("File tanda tangan pegawai tidak ditemukan.")
        if signature_anchor_paragraph is None:
            raise ValueError("Posisi tanda tangan pada template tidak ditemukan.")
        _add_floating_signature(signature_anchor_paragraph, signature)

    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            if "JejakASN" in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    "generate report by JejakASN", "Laporan dibuat otomatis oleh E-Master Jatim")

    activity_table = document.tables[1]
    _repeat_table_header(activity_table.rows[0])
    for cell in activity_table.rows[0].cells:
        _set_cell_text(cell, cell.text.strip(), WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    for old_row in list(activity_table.rows[1:]):
        activity_table._tbl.remove(old_row._tr)

    entries = activities or [ReportActivity(activity_time="-", text="Tidak ada aktivitas.")]
    for index, activity in enumerate(entries, start=1):
        row = activity_table.add_row()
        _set_cell_text(row.cells[0], str(index), WD_ALIGN_PARAGRAPH.CENTER)
        date_text = (f"{format_indonesian_date(report_date, include_day=True)}\n"
                     f"{activity.activity_time}")
        _set_cell_text(row.cells[1], date_text, WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_text(row.cells[2], activity.text, WD_ALIGN_PARAGRAPH.JUSTIFY)

        evidence_cell = row.cells[3]
        evidence_cell.text = ""
        evidence_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if not activity.evidence:
            _set_cell_text(evidence_cell, "-", WD_ALIGN_PARAGRAPH.CENTER)
        else:
            for evidence_index, evidence in enumerate(activity.evidence, start=1):
                picture_paragraph = evidence_cell.paragraphs[0] if evidence_index == 1 \
                    else evidence_cell.add_paragraph()
                _clear_paragraph(picture_paragraph)
                _format_paragraph(picture_paragraph, WD_ALIGN_PARAGRAPH.CENTER)
                if evidence.local_path and evidence.local_path.is_file():
                    _add_picture(picture_paragraph, evidence.local_path)
                link_paragraph = evidence_cell.add_paragraph()
                _format_paragraph(link_paragraph, WD_ALIGN_PARAGRAPH.CENTER)
                label = "Lihat File" if len(activity.evidence) == 1 else f"Lihat File {evidence_index}"
                _add_hyperlink(link_paragraph, label, evidence.url)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_run_font(run)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output
